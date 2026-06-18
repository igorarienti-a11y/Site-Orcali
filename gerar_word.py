from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Cores Orcali ──────────────────────────────────────────────────────────────
AZUL   = RGBColor(0x00, 0x2E, 0x6E)   # #002E6E
VERDE  = RGBColor(0x72, 0xBF, 0x45)   # #72BF45
CINZA  = RGBColor(0x44, 0x44, 0x44)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

# ── Margens ───────────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(text, level=1, color=AZUL, size=16, bold=True, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def eyebrow(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.font.size  = Pt(9)
    run.font.color.rgb = VERDE
    run.bold = True
    return p

def body(text, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = CINZA
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = CINZA
    return p

def numbered(text, size=10.5):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = CINZA
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '002E6E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def section_header(num, title, eyebrow_text):
    """Slide header block: número + eyebrow + título"""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f"SLIDE {num}")
    run.font.size  = Pt(8)
    run.font.color.rgb = VERDE
    run.bold = True

    eyebrow(eyebrow_text)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = AZUL
    divider()

def sub(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = AZUL
    return p

def table_simple(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, '002E6E')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = BRANCO
        run.font.size = Pt(10)
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(3)
    # data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'F0F4FB')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# CAPA DO DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("TRÁFEGO PAGO — ORCALI 2026")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = AZUL

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Para prospecções que não podem parar")
run2.font.size = Pt(13)
run2.font.color.rgb = VERDE
run2.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Conteúdo dos slides — referência para montagem no Canva")
run3.font.size = Pt(10)
run3.font.color.rgb = CINZA

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO: ANÁLISE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("▌ SEÇÃO 1 — ANÁLISE")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = AZUL

# ── SLIDE 1 ───────────────────────────────────────────────────────────────────
section_header(1, "Capa", "Abertura")
body("Headline: Tráfego Pago / Orcali")
body("Subheadline: Para prospecções que não podem parar")
body("→ Clonar slide 1 da social media (layout de capa com gradiente).")

# ── SLIDE 2 ───────────────────────────────────────────────────────────────────
section_header(2, "Objetivos de Performance", "Tráfego Pago")

sub("1 — Reconhecimento de marca nos nichos-alvo")
body("Construir presença da Orcali com decisores B2B em instituições de saúde, ensino e financeiras no PR e SC.")

sub("2 — Geração de leads qualificados")
body("Atrair gestores e diretores administrativos com intenção real de contratar serviços de facilities e segurança.")

sub("3 — Estrutura de nutrição e acompanhamento")
body("Conectar cada lead capturado ao RD Station para rastreamento da jornada, do primeiro clique ao fechamento.")
body("→ Clonar slide 2 da social media (3 objetivos numerados com linha divisória).")

# ── SLIDE 3 ───────────────────────────────────────────────────────────────────
section_header(3, "Diagnóstico", "Histórico 2024")

sub("O que foi feito")
body("~R$ 6.000/mês em tráfego pago ao longo de 2024.")

sub("O problema")
body('Resultado declarado: "baixo impacto comercial."')
body("Causa raiz: sem LP dedicada, sem rastreamento de origem, sem qualificação de lead, sem integração com CRM.")

sub("O que mudamos para 2026")
table_simple(
    ["Antes", "Agora"],
    [
        ["Site institucional genérico", "LP dedicada por nicho"],
        ["Sem rastreamento", "Pixel + CAPI + Google Tag"],
        ["Formulário básico", "Typebot com qualificação"],
        ["Lead perdido no e-mail", "Google Sheets + RD Station"],
        ["Sem atribuição", "UTM + origem em 38 campos"],
    ],
    [2.8, 2.8]
)

# ── SLIDE 4 ───────────────────────────────────────────────────────────────────
section_header(4, "Análise de Canais", "Meta vs Google para B2B")

sub("Por que Meta?")
body("O decisor B2B não levanta a mão facilmente. O Meta constrói reconhecimento antes da decisão de pesquisar — quando o gestor já conhece a Orcali, a busca no Google converte mais.")
bullet("Alcance segmentado por cargo, setor e região")
bullet("Retargeting de visitantes da LP")
bullet("Lookalike da base de 2.966 clientes ativos")

sub("Por que Google?")
body('70% da jornada B2B é autodidata. Quando o decisor pesquisa "empresa de facilities SC", ele já está próximo da decisão.')
bullet("Captura intenção direta de compra")
bullet("Keywords de fundo de funil por nicho")
bullet("Maior custo por lead, maior qualidade")

sub("A lógica da combinação")
body("Meta constrói familiaridade → Google captura a decisão.")
table_simple(
    ["", "Meta", "Google Search"],
    [
        ["Papel", "Alcance + Retargeting", "Fundo de funil"],
        ["Lead", "Mais frio", "Mais quente"],
        ["CPL estimado", "Menor", "R$ 80–R$ 150"],
        ["Quando converte", "Após nutrição", "Na busca ativa"],
    ],
    [1.5, 2.5, 2.5]
)

# ── SLIDE 5 ───────────────────────────────────────────────────────────────────
section_header(5, "Jornada do Decisor Online", "Como o B2B decide")

sub("Os dados (fontes: Think with Google, pesquisa B2B Brasil 2025)")
bullet("70% da jornada de compra B2B é feita sem contato com vendedor")
bullet("94% dos decisores estão informados antes de entrar em contato")
bullet("56% usam o Google pelo menos uma vez por semana na decisão")
bullet("86% já têm fornecedores em mente no início da jornada")
bullet("43% buscam novos fornecedores quando há insatisfação com o atual")
bullet("31% apontam confiança na marca como fator decisivo")

sub("O que isso significa para a Orcali")
body("O gestor não clica no primeiro anúncio e fecha. Ele pesquisa, compara, volta ao Google, visita o site — e aí decide. A presença contínua no Meta + a captura no Google formam a jornada completa.")

# ── SLIDE 6 ───────────────────────────────────────────────────────────────────
section_header(6, "Benchmarket Digital", "Concorrentes no ambiente pago")

sub("Orsegups — Tecnologia como diferencial")
bullet("Comunicação agressiva: IA, monitoramento inteligente, plataforma própria")
bullet("Certificações ISO 27001 e ISO 27701")
bullet("Escala nacional muito evidente")
bullet("Forte presença em busca orgânica (marca consolidada)")

sub("Orbenk — ESG e comunidade")
bullet("Campanha institucional ativa em SC (TV, rádio e portais digitais)")
bullet("Posicionamento: bem-estar, cuidado, transformação")
bullet("Tecnologia como inovação (robôs, drones, análise de dados)")
bullet("420 cidades, 33.000 colaboradores — escala como argumento")

sub("Khronos — Especialização em segurança eletrônica")
bullet("Monitoramento eletrônico como core")
bullet("Comunicação técnica e direta")

sub("Oportunidade Orcali")
body("Enquanto os concorrentes comunicam escala e tecnologia, a Orcali comunica 58 anos de continuidade, confiabilidade e presença regional — argumento mais relevante para o decisor B2B que tem medo de trocar de fornecedor.")

# ── SLIDE 7 ───────────────────────────────────────────────────────────────────
section_header(7, "Personas no Digital", "Como os decisores chegam até nós")

sub("Mariana Costa — Gerente Administrativa (Saúde)")
body('"Preciso garantir um ambiente seguro, limpo e eficiente para que as equipes médicas possam focar no atendimento."')
bullet('Busca no Google: "empresa facilities hospital SC", "terceirização limpeza hospitalar"')
bullet("Consulta LinkedIn e portais setoriais antes de entrar em contato")
bullet("Precisa de prova de conformidade e certificações antes de aprovar qualquer fornecedor")

sub("Ricardo Almeida — Diretor Administrativo (Ensino)")
body('"Quero que a operação funcione sem interferir na atividade educacional."')
bullet('Busca: "facilities instituição de ensino", "terceirização portaria escola"')
bullet("Pesquisa por cases e referências de outras instituições")
bullet("Decisão influenciada por diretoria — precisa de argumento para aprovar internamente")

sub("Comportamento em comum")
body("Chegam ao Meta passivamente (não buscam, são impactados). Chegam ao Google ativamente (já decidiram pesquisar). A conversão acontece após múltiplos pontos de contato.")

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO: ESTRATÉGIA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run("▌ SEÇÃO 2 — ESTRATÉGIA")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = AZUL

# ── SLIDE 8 ───────────────────────────────────────────────────────────────────
section_header(8, "Segmentação", "Nichos e Regiões")

sub("Nichos — Fase inicial (LP Facilities)")
table_simple(
    ["Nicho", "Por que priorizar"],
    [
        ["Instituições financeiras", "Alto ticket, decisão corporativa, necessidade de compliance"],
        ["Hospitais e clínicas", "Operação 24h, zero tolerância a falha, exige certificações"],
        ["Instituições de ensino", "Múltiplos prédios, gestão complexa, público recorrente"],
    ],
    [1.8, 3.8]
)

sub("Regiões — PR e SC")
bullet("SC (Always-on): Florianópolis, Joinville, Blumenau, Itajaí, Tubarão")
bullet("PR (Branding ativo): Curitiba, Toledo, Londrina, Guarapuava")
body("Geo alinhado ao briefing: SC = mercado base consolidado / PR = expansão.")

# ── SLIDE 9 ───────────────────────────────────────────────────────────────────
section_header(9, "Estratégia Meta", "Alcance + Retargeting")

sub("Campanha 1 — Reconhecimento de marca")
bullet("Objetivo: Alcance / Awareness")
bullet("Público 1: Lookalike 1% da base de 2.966 clientes ativos")
bullet("Público 2: Cargo (Gerente/Diretor Administrativo) + setor (saúde, educação, financeiro) + geo (SC/PR)")
bullet('Criativos: institucionais, "operações que não podem parar", cases por nicho')

sub("Campanha 2 — Retargeting")
bullet("Público: visitantes da LP (Pixel) + engajamento Instagram")
bullet("Objetivo: Geração de Leads / Conversão")
bullet("Criativos: prova social por nicho, diferenciais Orcali (58 anos, 1.554 clientes, certificações)")
bullet("Formato: carrossel com argumentos + CTA direto para o formulário")

sub("Lookalike")
body("Base: LEADS ORCALI.ods — 2.966 contatos. Normalizar e-mail + telefone antes do upload. Semente ideal: clientes com contratos ativos há mais de 5 anos.")

# ── SLIDE 10 ──────────────────────────────────────────────────────────────────
section_header(10, "Estratégia Google", "Fundo de Funil + Display")

sub("Search — Fundo de funil (prioridade)")
body("Grupo 1 — Saúde")
body('"empresa facilities hospital SC" · "terceirização limpeza hospitalar" · "segurança patrimonial hospital"', size=9.5)
body("Grupo 2 — Educação")
body('"facilities instituição de ensino" · "terceirização portaria escola" · "limpeza universidade SC"', size=9.5)
body("Grupo 3 — Financeiro")
body('"empresa facilities banco" · "terceirização serviços financeiro" · "segurança patrimonial Curitiba"', size=9.5)
body("Grupo 4 — Geral/Concorrente")
body('"Orcali facilities" · "alternativa Orbenk" · "empresa terceirização SC PR"', size=9.5)

sub("Display — Topo de funil")
bullet("Remarketing: visitantes da LP que não converteram")
bullet("Contextual: portais de notícia B2B, sites do setor de saúde/educação")
bullet("Criativo: banner com foco em continuidade operacional")

# ── SLIDE 11 ──────────────────────────────────────────────────────────────────
section_header(11, "Distribuição de Verba", "R$ 8.000 / mês")

table_simple(
    ["Canal", "Verba mensal", "%", "Objetivo"],
    [
        ["Google Ads (Search + Display)", "R$ 5.500", "69%", "Leads qualificados"],
        ["Meta Ads (Alcance + Retargeting)", "R$ 2.500", "31%", "Awareness + Retargeting"],
        ["TOTAL", "R$ 8.000", "100%", "—"],
    ],
    [2.5, 1.3, 0.7, 2.0]
)

sub("Racional")
body("B2B de alto ticket prioriza intenção de busca: o Google recebe a maior fatia porque o decisor chega com a decisão semi-formada. O Meta complementa com presença de marca e cobre o retargeting de quem visitou mas não converteu.")
body("Referência: alocação padrão B2B 2026 é 46% Google / 8% Meta (sem LinkedIn). Na Orcali, o Meta recebe mais por causa do retargeting e da base lookalike forte.")

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO: EXECUÇÃO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run("▌ SEÇÃO 3 — EXECUÇÃO")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = AZUL

# ── SLIDE 12 ──────────────────────────────────────────────────────────────────
section_header(12, "Landing Page + Formulário", "Estrutura de Captação")

sub("O que foi construído")
body("Landing page dedicada ao serviço de Facilities, desenvolvida especificamente para conversão:")
bullet("URL: site-orcali.vercel.app")
bullet("Estrutura: Hero → Diferenciais → Serviços → Prova social → Formulário")
bullet("Formulário qualificador via Typebot (chatbot integrado na página)")

sub("Campos do formulário")
body("Nome · E-mail corporativo · Telefone/WhatsApp · Empresa · Segmento · Estado · Cidade · Mensagem")

sub("Por que Typebot (e não formulário estático)?")
bullet("Taxa de conclusão maior (lógica conversacional)")
bullet("Qualificação por nicho/segmento antes de chegar ao comercial")
bullet("Rastreamento de UTMs automático desde o primeiro campo")
bullet("Integração direta com Sheets e CAPI em tempo real")

# ── SLIDE 13 ──────────────────────────────────────────────────────────────────
section_header(13, "Tracking + Integrações", "De onde vem o lead até o comercial")

sub("O fluxo completo")
body("Anúncio (Meta / Google)")
body("  ↓  LP — PageView → Meta CAPI + Google Tag")
body("  ↓  Typebot — form qualificador (UTMs + device + geo)")
body("  ↓  Edge Function (Vercel) — 38 campos gravados")
body("  ↓  Google Sheets — lead com origem, nicho, dispositivo, IP geo")
body("  ↓  Apps Script")
body("        ├── Meta CAPI (evento Lead — deduplica com Pixel)")
body("        └── RD Station (lead entra no funil de nutrição)")

sub("O que isso resolve")
body("Todo lead tem origem rastreada: de qual campanha veio, qual anúncio, qual nicho, qual cidade. O comercial recebe o lead com contexto — sabe se é hospital de Florianópolis que veio do Google ou banco de Curitiba que veio do retargeting Meta.")

# ── SLIDE 14 ──────────────────────────────────────────────────────────────────
section_header(14, "KPIs e Metas", "Como vamos medir")

table_simple(
    ["Indicador", "Meta (mês 1–2)", "Meta (mês 4–6)"],
    [
        ["CPL Google", "≤ R$ 150", "≤ R$ 100"],
        ["CPL Meta", "≤ R$ 90", "≤ R$ 60"],
        ["Leads qualificados/mês (Google)", "30–40", "50–60"],
        ["Taxa de conversão LP", "> 3%", "> 5%"],
        ["Taxa MQL (lead → oportunidade)", "8%", "12%"],
    ],
    [3.0, 1.8, 1.8]
)

sub("O que acompanhamos semanalmente")
bullet("Volume de leads por canal e por nicho")
bullet("CPL por campanha e por grupo de anúncio")
bullet("Qualidade: quantos leads viraram reunião agendada")
bullet("Score de qualidade dos anúncios (Google) e relevância (Meta)")

sub("O que muda ao longo do tempo")
body("Primeiros 60 dias: fase de aprendizado — CPL mais alto, é normal. A partir do mês 3: otimização com base em dados reais. Após 6 meses: CPL tende a cair 30–50%.")

# ── SLIDE 15 ──────────────────────────────────────────────────────────────────
section_header(15, "Próximos Passos", "O que acontece agora")

numbered("Finalizar integração RD Station (esta semana)\n→ Gerar token API em app.rdstation.com.br → Configurações → Integrações → Acessos à API")
numbered("Upload Lookalike Meta (semana 1)\n→ Normalizar base de 2.966 clientes (e-mail + telefone) → Custom Audience → Lookalike 1%")
numbered("Criar campanhas Meta (semana 1–2)\n→ Campanha de Alcance (Lookalike + cargo/setor) + Retargeting (visitantes LP)")
numbered("Criar campanhas Google (semana 1–2)\n→ Estrutura de grupos por nicho: saúde, ensino, financeiro + Display remarketing")
numbered("Criar criativos iniciais (semana 2)\n→ 3–4 criativos institucionais por canal, adaptados por nicho")
numbered("Primeiro relatório (dia 30)\n→ Análise de primeiros dados: CPL, volume, qualidade, ajustes")

# ══════════════════════════════════════════════════════════════════════════════
# RODAPÉ — FONTES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Referências e Fontes", level=2, size=12, color=CINZA, space_before=4)
divider()
body("Jornada do decisor B2B: Think with Google — thinkwithgoogle.com/intl/pt-br/tendencias-de-consumo/jornada-do-consumidor/b2b-jornada-pesquisa/", size=9)
body("Benchmarks Meta vs Google B2B: B2B PPC 2025 Report — thedigitalbloom.com | Meta ROAS: 51% / Google ROAS: 67%", size=9)
body("CPL Brasil serviços B2B: witu.digital/quanto-custa-google-ads — R$ 40–R$ 180 para serviços profissionais", size=9)
body("Concorrentes: grandesnomesdapropaganda.com.br (Orbenk) + gazetamercantil.digital", size=9)
body("Ad Library Meta: ads.facebook.com/ads/library — verificar manualmente: Orsegups, Orbenk, Khronos", size=9)

# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\igori\projetos\orcali-lp\Orcali-Trafego-2026.docx"
doc.save(out)
print(f"Salvo em: {out}")
