from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

AZUL  = RGBColor(0x00, 0x2E, 0x6E)
VERDE = RGBColor(0x72, 0xBF, 0x45)
CINZA = RGBColor(0x44, 0x44, 0x44)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(3)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.color.rgb = AZUL
        run.font.size = Pt(22)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = AZUL
        run.font.size = Pt(16)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = VERDE
        run.font.size = Pt(12)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = CINZA
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = CINZA
    return p

def numbered_item(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = CINZA
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, '002E6E')
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = BRANCO
        r.font.size = Pt(10)
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(3)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
            if ri % 2 == 0:
                set_cell_bg(cell, 'F0F4FB')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TÍTULO DO DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Tráfego Pago — Orcali 2026")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = AZUL

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("Para prospecções que não podem parar")
run2.font.size = Pt(13)
run2.font.color.rgb = VERDE
run2.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(16)
run3 = p3.add_run("Planejamento estratégico de mídia paga — Meta Ads + Google Ads\nDigitha para Orcali Segurança e Serviços · Junho 2026")
run3.font.size = Pt(10)
run3.font.color.rgb = CINZA

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 — ANÁLISE
# ══════════════════════════════════════════════════════════════════════════════
h1("Análise")

# ── Objetivos de Performance ──────────────────────────────────────────────────
h2("Objetivos de Performance")

h3("1. Reconhecimento de marca nos nichos-alvo")
body("Construir presença contínua da Orcali junto a decisores B2B em instituições de saúde, ensino e financeiras no PR e SC, antes que eles comecem a pesquisar ativamente por fornecedores.")

h3("2. Geração de leads qualificados")
body("Atrair gerentes e diretores administrativos com intenção real de contratar serviços de facilities e segurança, com rastreamento de origem desde o primeiro clique.")

h3("3. Estrutura de nutrição e acompanhamento")
body("Conectar cada lead capturado ao RD Station para rastreamento completo da jornada comercial — do primeiro contato ao fechamento de contrato.")

# ── Diagnóstico ───────────────────────────────────────────────────────────────
h2("Diagnóstico — O que aconteceu em 2024")

h3("O investimento anterior")
body("A Orcali rodou aproximadamente R$ 6.000/mês em tráfego pago ao longo de 2024. O resultado foi descrito internamente como de baixo impacto comercial.")

h3("Por que não funcionou")
bullet("Sem landing page dedicada — o tráfego ia para o site institucional genérico")
bullet("Sem rastreamento de origem — impossível saber qual campanha gerava resultado")
bullet("Sem qualificação de lead — qualquer contato era tratado como igual")
bullet("Sem integração com CRM — lead chegava por e-mail e se perdia")

h3("O que mudamos para 2026")
table(
    ["Antes", "Agora"],
    [
        ["Site institucional genérico", "Landing page dedicada por nicho"],
        ["Sem rastreamento", "Meta Pixel + CAPI + Google Tag"],
        ["Formulário básico", "Typebot com qualificação por segmento"],
        ["Lead perdido no e-mail", "Google Sheets + RD Station integrados"],
        ["Sem atribuição de origem", "UTM + 38 campos de dados por lead"],
    ],
    [2.8, 2.8]
)

# ── Análise de canais ─────────────────────────────────────────────────────────
h2("Análise de Canais — Meta vs Google para B2B")

h3("Por que Meta Ads")
body("O decisor B2B não levanta a mão facilmente. O Meta constrói reconhecimento antes da decisão de pesquisar — quando o gestor já conhece a Orcali, a busca no Google converte com muito mais facilidade.")
bullet("Alcance segmentado por cargo, setor e região")
bullet("Retargeting de visitantes da landing page")
bullet("Lookalike da base de 2.966 clientes ativos da Orcali")

h3("Por que Google Ads")
body("70% da jornada B2B é autodidata. Quando o decisor digita no Google, ele já está próximo da decisão — é o momento de capturar.")
bullet("Captura intenção direta de compra no momento da busca")
bullet("Keywords de fundo de funil segmentadas por nicho")
bullet("Custo por lead mais alto, mas qualidade significativamente superior")

h3("A lógica da combinação")
body("Meta constrói familiaridade → Google captura a decisão.")
table(
    ["", "Meta Ads", "Google Search"],
    [
        ["Papel na estratégia", "Alcance + Retargeting", "Fundo de funil"],
        ["Temperatura do lead", "Mais frio", "Mais quente"],
        ["CPL estimado", "R$ 50–90", "R$ 80–150"],
        ["Quando converte", "Após múltiplos contatos", "Na busca ativa"],
        ["MQL → SQL", "5–10%", "7–12%"],
    ],
    [2.2, 2.0, 2.0]
)

# ── Jornada do decisor ────────────────────────────────────────────────────────
h2("Jornada do Decisor Online")

h3("Como o B2B decide — dados de mercado")
bullet("70% da jornada de compra B2B é feita sem contato com vendedor (pesquisa autodidata)")
bullet("94% dos decisores já estão informados antes de entrar em contato com uma empresa")
bullet("56% usam o Google pelo menos uma vez por semana durante o processo de decisão")
bullet("86% já têm fornecedores em mente quando começam a jornada de compra")
bullet("43% buscam novos fornecedores quando há insatisfação com o atual")
bullet("31% apontam confiança na marca como fator número 1 na contratação B2B")
body("Fontes: Think with Google — Jornada de Compra B2B | Pesquisa B2B Brasil 2025")

h3("O que isso significa para a Orcali")
body("O gestor hospitalar ou diretor escolar não clica no primeiro anúncio e fecha contrato. Ele pesquisa passivamente, compara, volta ao Google semanas depois, visita o site, vê referências — e aí decide. A presença contínua no Meta constrói o terreno para a conversão no Google.")

# ── Benchmarket digital ───────────────────────────────────────────────────────
h2("Benchmarket Digital — Concorrentes")

h3("Orsegups")
bullet("Posicionamento forte em tecnologia e inovação em segurança")
bullet("Comunicação agressiva: IA, monitoramento inteligente, plataforma própria, menor tempo de resposta")
bullet("Certificações de segurança da informação: ISO 27001 e ISO 27701")
bullet("Escala nacional muito evidente na comunicação")
bullet("Presença consolidada em busca orgânica — palavra 'segurança' associada à marca")

h3("Orbenk")
bullet("Campanha institucional ativa em SC — TV, rádio e portais digitais")
bullet("Posicionamento: bem-estar, cuidado, transformação, ESG")
bullet("Tecnologia como inovação: robôs de limpeza, drones, análise de dados")
bullet("420 cidades, 33.000 colaboradores — escala como argumento de credibilidade")

h3("Khronos")
bullet("Especialização em monitoramento eletrônico")
bullet("Comunicação técnica focada em soluções integradas de segurança")

h3("Oportunidade para a Orcali")
body("Enquanto os concorrentes comunicam escala e tecnologia, a Orcali tem um argumento que eles não têm: 58 anos de continuidade, confiabilidade comprovada e presença regional profunda. Para o decisor B2B que teme trocar de fornecedor e se arrepender, esse é o argumento mais relevante.")

# ── Personas no digital ───────────────────────────────────────────────────────
h2("Personas no Digital")

h3("Mariana Costa — Gerente Administrativa (Saúde, 42 anos)")
body('"Preciso garantir um ambiente seguro, limpo e eficiente para que as equipes médicas possam focar no atendimento."')
bullet("Busca no Google: empresa facilities hospital SC, terceirização limpeza hospitalar, segurança patrimonial hospital Florianópolis")
bullet("Consulta LinkedIn e portais setoriais antes de entrar em contato")
bullet("Exige prova de certificações e conformidade regulatória")
bullet("Responde diretamente à diretoria executiva — precisa de dados para justificar a escolha")

h3("Ricardo Almeida — Diretor Administrativo (Ensino, 42 anos)")
body('"Quero que a operação funcione sem interferir na atividade educacional."')
bullet("Busca: facilities instituição de ensino, terceirização portaria escola, empresa limpeza universidade")
bullet("Pesquisa por cases e referências de outras instituições semelhantes")
bullet("Precisa convencer a diretoria — o argumento de custo-benefício é central")
bullet("Sensível a reputação da instituição e experiência de alunos e visitantes")

h3("Comportamento digital em comum")
body("Chegam ao Meta passivamente — não estão procurando, mas são impactados pela Orcali enquanto navegam. Chegam ao Google ativamente — já tomaram a decisão de pesquisar. A conversão acontece após múltiplos pontos de contato entre as duas plataformas.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 — ESTRATÉGIA
# ══════════════════════════════════════════════════════════════════════════════
h1("Estratégia")

# ── Nichos + Geo ──────────────────────────────────────────────────────────────
h2("Segmentação — Nichos e Regiões")

h3("Nichos prioritários — fase inicial")
table(
    ["Nicho", "Por que priorizar"],
    [
        ["Instituições financeiras", "Alto ticket, decisão corporativa, compliance obrigatório"],
        ["Hospitais e clínicas", "Operação 24h, zero tolerância a falha, exige ISO e certificações"],
        ["Instituições de ensino", "Múltiplos prédios, gestão complexa, ciclo de renovação previsível"],
    ],
    [2.2, 3.4]
)

h3("Regiões de atuação")
bullet("Santa Catarina (always-on): Florianópolis, Joinville, Blumenau, Itajaí, Tubarão")
bullet("Paraná (expansão ativa): Curitiba, Toledo, Londrina, Guarapuava")
body("Lógica geográfica: SC é o mercado base consolidado onde a Orcali já tem presença e pode converter mais rápido. PR é expansão onde o awareness precisa ser construído primeiro.")

# ── Estratégia Meta ───────────────────────────────────────────────────────────
h2("Estratégia Meta Ads")

h3("Campanha 1 — Reconhecimento de Marca")
bullet("Objetivo de campanha: Alcance / Awareness")
bullet("Público A: Lookalike 1% gerado a partir da base de 2.966 clientes ativos da Orcali")
bullet("Público B: Segmentação por cargo (Gerente/Diretor Administrativo) + setor (saúde, educação, financeiro) + geo PR e SC")
bullet("Criativos: institucionais, diferenciais da Orcali, cases por nicho, tom consultivo e confiável")

h3("Campanha 2 — Retargeting")
bullet("Público: visitantes da landing page rastreados pelo Pixel + engajamento no Instagram")
bullet("Objetivo de campanha: Geração de Leads / Conversão")
bullet("Criativos: prova social por nicho, diferenciais específicos (58 anos, 1.554 clientes, certificações)")
bullet("Formato principal: carrossel com argumentos por dor + CTA direto para formulário")

h3("Lookalike — base de clientes Orcali")
body("A Orcali tem uma base de 2.966 clientes cadastrados com e-mail e telefone. Essa base será carregada no Meta como Custom Audience para gerar um Lookalike 1% — o público mais semelhante aos clientes reais da Orcali dentro das regiões de SC e PR.")

# ── Estratégia Google ─────────────────────────────────────────────────────────
h2("Estratégia Google Ads")

h3("Search — Fundo de funil (prioridade principal)")
body("Estrutura por grupos de anúncio, segmentados por nicho:")
table(
    ["Grupo", "Palavras-chave principais"],
    [
        ["Saúde", "empresa facilities hospital SC, terceirização limpeza hospitalar, segurança patrimonial hospital"],
        ["Educação", "facilities instituição de ensino, terceirização portaria escola, limpeza universidade SC"],
        ["Financeiro", "empresa facilities banco, terceirização serviços financeiro, segurança patrimonial Curitiba"],
        ["Geral / Concorrente", "Orcali facilities, alternativa Orbenk, empresa terceirização SC PR"],
    ],
    [1.5, 4.1]
)

h3("Display — Topo de funil")
bullet("Remarketing: visitantes da landing page que não converteram no formulário")
bullet("Contextual: portais de notícias B2B, sites do setor de saúde e educação")
bullet("Criativo: banners com foco em continuidade operacional e confiabilidade")

# ── Distribuição de verba ─────────────────────────────────────────────────────
h2("Distribuição de Verba — R$ 8.000/mês")

table(
    ["Canal", "Verba mensal", "Participação", "Objetivo principal"],
    [
        ["Google Ads (Search + Display)", "R$ 5.500", "69%", "Leads qualificados — fundo de funil"],
        ["Meta Ads (Alcance + Retargeting)", "R$ 2.500", "31%", "Awareness + Retargeting"],
        ["Total", "R$ 8.000", "100%", "—"],
    ],
    [2.4, 1.2, 1.1, 2.0]
)

h3("Racional da distribuição")
body("B2B de alto ticket e ciclo de venda longo prioriza intenção de busca. O Google recebe a maior fatia porque o decisor chega com a decisão quase formada — o custo por lead é maior, mas a taxa de conversão compensa. O Meta complementa com presença de marca antes da busca e cobre o retargeting de quem visitou mas não converteu.")
body("Referência de mercado: a alocação padrão para B2B em 2026 é 46% Google + 8% Meta (sem LinkedIn). Na Orcali, o Meta recebe proporção maior por causa da base lookalike forte e da necessidade de construção de marca regional.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 — EXECUÇÃO
# ══════════════════════════════════════════════════════════════════════════════
h1("Execução")

# ── LP + Form ─────────────────────────────────────────────────────────────────
h2("Landing Page e Formulário")

h3("O que foi construído")
body("Desenvolvemos uma landing page dedicada ao serviço de Facilities, projetada especificamente para conversão — sem o ruído do site institucional. A página usa o posicionamento estratégico da Orcali e direciona o visitante direto para o formulário qualificador.")
bullet("URL: site-orcali.vercel.app")
bullet("Estrutura: Hero → Diferenciais → Serviços → Prova social → Formulário")
bullet("Formulário conversacional via Typebot (chatbot embutido na página)")

h3("Campos do formulário")
body("Nome · E-mail corporativo · Telefone/WhatsApp · Empresa · Segmento (saúde, ensino, financeiro, outros) · Estado · Cidade · Mensagem")

h3("Por que Typebot em vez de formulário estático")
bullet("Taxa de conclusão maior — a lógica conversacional reduz abandono")
bullet("Qualificação automática por nicho/segmento antes de chegar ao comercial")
bullet("Captura de UTMs desde o primeiro campo — origem rastreada sem falha")
bullet("Integração em tempo real com Google Sheets e Meta CAPI")

# ── Tracking + Integrações ────────────────────────────────────────────────────
h2("Tracking e Integrações")

h3("Fluxo completo do lead")
numbered_item("Usuário clica no anúncio (Meta ou Google) — UTMs registradas")
numbered_item("Landing page carrega — PageView enviado ao Meta CAPI e Google Tag")
numbered_item("Usuário preenche o Typebot — device, geo, IP, UTMs e dados do form capturados")
numbered_item("Edge function (Vercel) grava 38 campos no Google Sheets em tempo real")
numbered_item("Apps Script dispara: evento Lead ao Meta CAPI + lead criado no RD Station")
numbered_item("Comercial recebe lead com origem completa — campanha, nicho, cidade, dispositivo")

h3("O que isso resolve na prática")
body("Todo lead chega ao comercial com contexto: veio do Google buscando 'facilities hospital SC' ou do retargeting Meta após visitar a página? É gerente de hospital em Florianópolis ou diretor de escola em Curitiba? O comercial aborda com informação, não no escuro.")

h3("Stack técnica implantada")
bullet("Meta Pixel ID: 1650240729249255 — eventos PageView e Lead")
bullet("Meta CAPI — deduplicação automática com o Pixel")
bullet("Google Tag — rastreamento de conversão")
bullet("RD Station — pendente token API do cliente")

# ── KPIs ──────────────────────────────────────────────────────────────────────
h2("KPIs e Metas")

h3("Indicadores por fase")
table(
    ["Indicador", "Meta mês 1–2", "Meta mês 4–6"],
    [
        ["CPL Google Ads", "≤ R$ 150", "≤ R$ 100"],
        ["CPL Meta Ads", "≤ R$ 90", "≤ R$ 60"],
        ["Leads qualificados/mês — Google", "30–40", "50–60"],
        ["Taxa de conversão da landing page", "> 3%", "> 5%"],
        ["Taxa MQL (lead → oportunidade)", "8%", "12%"],
    ],
    [2.8, 1.5, 1.5]
)

h3("Referências de mercado usadas")
bullet("Google Ads — serviços profissionais B2B Brasil: R$ 40–R$ 180 por lead")
bullet("Meta Ads — B2B: CPL menor, taxa de qualificação de 5–10% MQL para SQL")
bullet("Google Search — B2B: taxa de qualificação de 7–12% MQL para SQL")
bullet("Após 6 meses de otimização contínua: CPL tende a cair 30–50%")

h3("Acompanhamento semanal")
bullet("Volume de leads por canal e por nicho")
bullet("CPL por campanha e por grupo de anúncio")
bullet("Qualidade: leads que viraram reunião agendada")
bullet("Score de qualidade (Google) e relevância (Meta)")

# ── Próximos passos ───────────────────────────────────────────────────────────
h2("Próximos Passos")

h3("Semana 1")
numbered_item("Finalizar integração RD Station — cliente gera token em app.rdstation.com.br → Configurações → Integrações → Acessos à API")
numbered_item("Normalizar base de 2.966 clientes (e-mail + telefone) e fazer upload no Meta para criar Custom Audience e Lookalike 1%")

h3("Semanas 1–2")
numbered_item("Criar campanhas Meta: Alcance (Lookalike + cargo/setor/geo) + Retargeting (visitantes da LP)")
numbered_item("Criar campanhas Google: grupos por nicho (saúde, ensino, financeiro) + Display remarketing")
numbered_item("Produzir criativos iniciais — 3 a 4 peças por canal, adaptadas por nicho")

h3("Dia 30 — Primeiro relatório")
numbered_item("Análise de primeiros dados: CPL real por canal, volume de leads, qualidade, ajustes de segmentação e criativos")

# ══════════════════════════════════════════════════════════════════════════════
# FONTES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h2("Referências e Fontes")
body("Jornada do decisor B2B: Think with Google — thinkwithgoogle.com/intl/pt-br/tendencias-de-consumo/jornada-do-consumidor/b2b-jornada-pesquisa/")
body("Benchmarks plataformas B2B: B2B PPC 2025 Report — thedigitalbloom.com")
body("Meta B2B ROAS 51% / Google Search ROAS 67% — Lever Digital, 2026")
body("CPL Brasil serviços profissionais B2B: witu.digital/quanto-custa-google-ads")
body("Concorrentes: grandesnomesdapropaganda.com.br (Orbenk) | gazetamercantil.digital")
body("Ad Library Meta — verificar manualmente: ads.facebook.com/ads/library (buscar: Orsegups, Orbenk, Khronos)")

# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\igori\projetos\orcali-lp\Orcali-Trafego-Gamma.docx"
doc.save(out)
print(f"Salvo em: {out}")
